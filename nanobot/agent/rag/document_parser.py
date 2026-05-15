"""RAG文档解析器 - 多格式文档内容提取"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from loguru import logger


@dataclass
class ParsedDocument:
    """解析后的文档"""

    content: str
    doc_type: str
    file_path: str
    metadata: dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentParser:
    """多格式文档解析器

    支持格式：
    - PDF → MarkItDown → Markdown文本
    - DOCX/DOC → python-docx → 纯文本
    - Markdown/TXT → 直接读取
    - 代码文件 → 保留结构信息
    """

    _PARSER_MAP: dict[str, str] = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".md": "text",
        ".markdown": "text",
        ".mdx": "text",
        ".txt": "text",
        ".rst": "text",
        ".adoc": "text",
        ".py": "code",
        ".js": "code",
        ".ts": "code",
        ".jsx": "code",
        ".tsx": "code",
        ".java": "code",
        ".go": "code",
        ".rs": "code",
        ".c": "code",
        ".cpp": "code",
        ".h": "code",
        ".hpp": "code",
        ".cs": "code",
        ".rb": "code",
        ".php": "code",
        ".swift": "code",
        ".kt": "code",
        ".scala": "code",
        ".rust": "code",
        ".html": "markup",
        ".htm": "markup",
        ".xml": "markup",
        ".svg": "markup",
        ".json": "data",
        ".yaml": "data",
        ".yml": "data",
        ".toml": "data",
        ".ini": "data",
        ".cfg": "data",
        ".csv": "data",
        ".tsv": "data",
        ".log": "text",
        ".sh": "code",
        ".bash": "code",
        ".zsh": "code",
        ".ps1": "code",
        ".bat": "code",
        ".sql": "code",
    }

    @classmethod
    def detect_type(cls, file_path: str | Path) -> str:
        """根据扩展名检测文档类型"""
        ext = Path(file_path).suffix.lower()
        return cls._PARSER_MAP.get(ext, "text")

    @classmethod
    def parse(cls, file_path: str | Path, doc_type: str | None = None) -> ParsedDocument:
        """解析文件，返回统一格式的ParsedDocument"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not path.is_file():
            raise ValueError(f"Not a regular file: {file_path}")

        detected_type = doc_type or cls.detect_type(path)

        parser_method = getattr(cls, f"_parse_{detected_type}", cls._parse_text)
        content = parser_method(str(path))

        return ParsedDocument(
            content=content,
            doc_type=detected_type,
            file_path=str(path),
            metadata={
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "extension": path.suffix.lower(),
            },
        )

    @staticmethod
    def _read_text_with_fallback(file_path: str) -> str:
        """尝试多种编码读取文本文件"""
        path = Path(file_path)
        encodings = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]
        for enc in encodings:
            try:
                return path.read_text(encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError(f"Cannot decode file with any supported encoding: {file_path}")

    @staticmethod
    def _parse_text(file_path: str) -> str:
        """纯文本 / Markdown / 通用文本"""
        return DocumentParser._read_text_with_fallback(file_path)

    @staticmethod
    def _parse_code(file_path: str) -> str:
        """代码文件 - 原样返回（分块由chunker处理）"""
        return DocumentParser._read_text_with_fallback(file_path)

    @staticmethod
    def _parse_markup(file_path: str) -> str:
        """HTML/XML等标记语言"""
        return DocumentParser._read_text_with_fallback(file_path)

    @staticmethod
    def _parse_data(file_path: str) -> str:
        """数据文件（JSON/YAML/CSV等）"""
        return DocumentParser._read_text_with_fallback(file_path)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """PDF解析 - 使用MarkItDown"""
        try:
            from markitdown import MarkItDown

            md = MarkItDown()
            result = md.convert(file_path)
            text = result.text_content

            if not text or not text.strip():
                logger.warning(f"MarkItDown returned empty content for PDF: {file_path}")

            return text

        except ImportError:
            raise ValueError(
                "PDF parsing requires markitdown. "
                "Install with: pip install markitdown"
            )
        except Exception as e:
            logger.error(f"MarkItDown failed to parse PDF {file_path}: {e}")
            raise ValueError(f"PDF parse error: {e}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """DOCX/DOC解析 - 使用python-docx"""
        try:
            from docx import Document

            doc = Document(file_path)

            parts: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style_name = para.style.name if para.style else ""
                    if style_name.startswith("Heading"):
                        level = style_name.replace("Heading", "").strip() or "1"
                        parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
                    else:
                        parts.append(text)

            if doc.tables:
                parts.append("\n## Tables\n")
                for table_idx, table in enumerate(doc.tables, 1):
                    parts.append(f"### Table {table_idx}")
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append(" | ".join(cells))
                    parts.append("")

            return "\n".join(parts)

        except ImportError:
            raise ValueError(
                "DOCX parsing requires python-docx. "
                "Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise ValueError(f"DOCX parse error: {e}")

    @staticmethod
    def _parse_doc(file_path: str) -> str:
        """旧版 .doc 格式 - 尝试转换为文本"""
        try:
            import subprocess

            result = subprocess.run(
                ["catdoc", "-w", file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        try:
            import olefile

            ole = olefile.OleFileIO(file_path)
            stream = ole.openstream("WordDocument")
            raw = stream.read()
            ole.close()

            decoded = raw.decode("utf-8", errors="ignore")
            text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", decoded)
            return text.strip()

        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"OLE parsing for .doc failed: {e}")

        raise ValueError(
            "Legacy .doc format support is limited. "
            "Consider converting to .docx first."
        )
